import streamlit as st
import openpyxl
import random
import re
import io
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PAGE CONFIGURATION ---
st.set_page_config(page_title="PE Exam Simulator", page_icon="📝", layout="wide")

# --- STUDY PLAN DATA (Inserted for Targeted Drills) ---
STUDY_PLAN_DATA = {
    "Fire Dynamics": [
        "MF-69", "MF-124", "MF-25", "MF-66", "MF-163", "MF-14",
    ],
    "Hydraulics & Pumps": [
        "MF-170", "MISC-43", "MF-32", "MF-44", "MF-41", "SoPE-40",
        "MISC-53", "MISC-23", "MF-168", "MF-109", "MISC-59", "MF-36"
    ],
    "Special Hazards": [
        "SoPE-48", "MF-45", "SoPE-49", "MF-131", "MISC-66", "GEN_69_292",
        "MF-136", "MISC-57", "SoPE-64", "MISC-63", "MF-62", "SoPE-78"
    ],
    "Risk & Reliability": [
        "MF-4", "MF-3", "MF-160", "MF-150", "MISC-70"
    ],
    "Fire Alarm & Detection": [
        "MISC-45", "MISC-11", "MF-55", "MF-143", "MF-54", "MF-58"
    ],
    "Life Safety & Egress": [
        "MF-146", "MISC-58", "MF-92", "MF-176", "SoPE-81", "MF-82",
        "MF-77", "MF-75"
    ]
}


# --- CLASS DEFINITION ---
class Question:
    def __init__(self, text, options, correct_indices, q_id):
        self.text = text
        self.options = options
        self.correct_indices = set(correct_indices)
        self.id = q_id
        self.user_selections = set()
        self.flagged = False

    def is_correct(self):
        return self.user_selections == self.correct_indices

    def is_answered(self):
        return len(self.user_selections) > 0


# --- HELPER FUNCTIONS ---
def parse_excel(uploaded_file):
    """Parses the uploaded Excel file into Question objects."""
    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    qs = []

    for sheet in wb.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=False), start=2):
            if not row[0].value: continue
            txt = str(row[0].value)

            qid_match = re.search(r'\(([\w-]+)\)\s*$', txt)
            if qid_match:
                qid = qid_match.group(1)
            else:
                qid = f"GEN_{row_idx}_{random.randint(100, 999)}"

            opts = []
            cor = []
            for i, c in enumerate(row[1:]):
                if c.value:
                    opts.append(str(c.value))
                    if c.font and c.font.bold:
                        cor.append(i)

            if opts and cor:
                qs.append(Question(txt, opts, cor, qid))
    return qs


def generate_word_doc(questions):
    """Generates a Word document in memory."""
    doc = Document()
    doc.add_heading('PE Exam Set', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated ID: {random.randint(1000, 9999)}")

    for i, q in enumerate(questions):
        p = doc.add_paragraph()
        p.add_run(f"{i + 1}. {q.text}").bold = True
        lets = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        for idx, o in enumerate(q.options):
            doc.add_paragraph(f"    {lets[idx] if idx < 26 else '?'}. {o}")
        doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading('Answer Key', 0)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Shading'
    tbl.rows[0].cells[0].text = 'Q#'
    tbl.rows[0].cells[1].text = 'Ans'

    for i, q in enumerate(questions):
        row = tbl.add_row().cells
        row[0].text = str(i + 1)
        lets = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        ans_str = ", ".join([lets[x] for x in sorted(q.correct_indices)])
        row[1].text = ans_str

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_war_report(questions, title="Wrong Answer Roundup"):
    """Generates a text file listing wrong questions for review."""
    output = f"--- {title} ---\n\n"
    for q in questions:
        if not q.is_correct():
            output += f"ID: {q.id}\n"
            output += f"Q: {q.text}\n"
            lets = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
            ans_str = ", ".join([lets[x] for x in sorted(q.correct_indices)])
            output += f"Correct Answer: {ans_str}\n"
            output += "-" * 40 + "\n"
    return output


# --- SESSION STATE INITIALIZATION ---
if 'questions_pool' not in st.session_state:
    st.session_state.questions_pool = []
if 'exam_session' not in st.session_state:
    st.session_state.exam_session = []
if 'current_index' not in st.session_state:
    st.session_state.current_index = 0
if 'exam_stage' not in st.session_state:
    st.session_state.exam_stage = "SETUP"

# Persistent Tracking Variables
if 'used_ids' not in st.session_state:
    st.session_state.used_ids = set()
if 'wrong_ids' not in st.session_state:
    st.session_state.wrong_ids = set()
if 'cumulative_correct' not in st.session_state:
    st.session_state.cumulative_correct = 0
if 'cumulative_answered' not in st.session_state:
    st.session_state.cumulative_answered = 0


# --- CALLBACKS ---
def load_history_callback():
    """Run this ONLY when a new history file is uploaded."""
    uploaded = st.session_state.history_uploader
    if uploaded is not None:
        try:
            data = json.load(uploaded)
            if isinstance(data, list):
                st.session_state.used_ids.update(set(data))
                st.toast("Loaded legacy history (Progress only).")
            elif isinstance(data, dict):
                st.session_state.used_ids.update(set(data.get("used_ids", [])))
                st.session_state.wrong_ids.update(set(data.get("wrong_ids", [])))
                st.session_state.cumulative_correct = data.get("correct", 0)
                st.session_state.cumulative_answered = data.get("answered", 0)
                st.toast(f"History Restored.")
        except Exception as e:
            st.error(f"Error loading history: {e}")


# --- MAIN APP LOGIC ---

# 1. SIDEBAR: Controls & Navigation
with st.sidebar:
    st.header("1. Load Data")

    # A. Load Questions
    uploaded_file = st.file_uploader("Upload Question Database (Excel)", type=["xlsx"])
    if uploaded_file:
        try:
            if not st.session_state.questions_pool:
                questions = parse_excel(uploaded_file)
                st.session_state.questions_pool = questions
                st.success(f"Loaded {len(questions)} questions.")
        except Exception as e:
            st.error(f"Error parsing file: {e}")

    # B. Load History
    st.file_uploader(
        "Upload History File (Optional)",
        type=["json"],
        key="history_uploader",
        on_change=load_history_callback
    )

    # C. Performance Metrics (Simplified)
    st.write("---")
    st.header("📊 Performance")

    total_bank = len(st.session_state.questions_pool)
    used_count = len(st.session_state.used_ids)

    # Bank Progress ONLY
    st.caption("Bank Completion")
    prog_val = used_count / total_bank if total_bank > 0 else 0
    st.progress(prog_val)
    st.write(f"**{used_count} / {total_bank}** Questions Seen")

    # D. Download Wrong IDs (Anytime)
    if len(st.session_state.wrong_ids) > 0:
        st.write("---")
        wrong_ids_text = "All-Time Wrong Question IDs:\n" + "\n".join(sorted(list(st.session_state.wrong_ids)))
        st.download_button(
            label="⚠️ Download All Wrong IDs",
            data=wrong_ids_text,
            file_name="All_Wrong_IDs.txt",
            mime="text/plain"
        )

    # E. Exam Controls
    if st.session_state.questions_pool and st.session_state.exam_stage == "SETUP":
        st.write("---")
        st.header("2. Start Exam")

        # --- MODE SELECTION LOGIC ---
        mode = st.radio("Choose Mode:", ["Standard Exam (Random 20)", "Targeted Weakness Drill"])

        target_pool = []
        btn_label = "Start Exam"

        if mode == "Standard Exam (Random 20)":
            # Standard Logic: Filter out used questions
            target_pool = [q for q in st.session_state.questions_pool if q.id not in st.session_state.used_ids]

            if len(target_pool) == 0:
                st.warning("All questions attempted!")
                if st.button("🔄 Reset History & Scores", type="primary"):
                    st.session_state.used_ids = set()
                    st.session_state.wrong_ids = set()
                    st.session_state.cumulative_correct = 0
                    st.session_state.cumulative_answered = 0
                    st.rerun()
            else:
                remaining_qs = len(target_pool)
                btn_label = "Start New Standard Block"
                if remaining_qs < 20:
                    st.info(f"Only {remaining_qs} questions remaining.")
                    btn_label = f"Start Final {remaining_qs} Questions"

        else:  # Targeted Mode
            focus_area = st.selectbox("Select Area to Drill:", list(STUDY_PLAN_DATA.keys()))
            target_ids = set(STUDY_PLAN_DATA[focus_area])

            # Targeted Logic: Search ENTIRE pool (even used ones) for these IDs
            target_pool = [q for q in st.session_state.questions_pool if q.id in target_ids]

            st.info(f"Found {len(target_pool)} questions for **{focus_area}**.")
            btn_label = f"Start {focus_area} Drill"

        # --- START BUTTON LOGIC ---
        if len(target_pool) > 0:
            if st.button(btn_label, type="primary"):
                # If Targeted, we take ALL found questions (up to 20 or more? Let's cap at 20 for sanity)
                session_size = min(20, len(target_pool))

                # If Standard, sample random. If Targeted, also sample random from the filtered set?
                # Yes, random sample ensures variety if the set is > 20.
                st.session_state.exam_session = random.sample(target_pool, session_size)

                # Mark as used immediately upon start (Standard Mode behavior)
                # Note: In Targeted Mode, this doesn't hurt, but won't stop them from appearing again in Targeted Mode.
                for q in st.session_state.exam_session:
                    st.session_state.used_ids.add(q.id)
                    q.user_selections = set()
                    q.flagged = False

                st.session_state.current_index = 0
                st.session_state.exam_stage = "ACTIVE"
                st.rerun()

    # F. Save Progress
    st.write("---")
    st.header("3. Save Progress")
    if len(st.session_state.used_ids) > 0:
        save_data = {
            "used_ids": list(st.session_state.used_ids),
            "wrong_ids": list(st.session_state.wrong_ids),
            "correct": st.session_state.cumulative_correct,
            "answered": st.session_state.cumulative_answered
        }

        history_json = json.dumps(save_data)
        st.download_button(
            label="💾 Download Progress File",
            data=history_json,
            file_name="PE_Exam_History.json",
            mime="application/json"
        )

    # Reset Button
    if st.session_state.exam_stage != "SETUP":
        st.write("---")
        if st.button("End Exam / Return to Menu"):
            st.session_state.exam_stage = "SETUP"
            st.rerun()

    # Word Doc Export
    if st.session_state.questions_pool and st.session_state.exam_stage == "SETUP":
        if st.button("Generate Word Doc"):
            subset = random.sample(st.session_state.questions_pool, min(20, len(st.session_state.questions_pool)))
            docx_file = generate_word_doc(subset)
            st.download_button(
                label="Download .docx",
                data=docx_file,
                file_name="PE_Exam_Set.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# 2. MAIN CONTENT AREA
if st.session_state.exam_stage == "SETUP":
    st.title("PE Exam Simulator & Generator")
    st.info("""
    **Campaign Mode Active**
    1. **Upload** Excel Question Bank & History File.
    2. **Take Exams** (Questions are removed from the pool as you use them).
    3. **Download Reports:** - Get a "Session Report" immediately after an exam.
       - Get an "All-Time Wrong IDs" list from the sidebar anytime.
    4. **Save Progress** before leaving.
    """)

    # Show active mode help text
    st.markdown("### How to use Targeted Drills:")
    st.markdown("""
    - Select **'Targeted Weakness Drill'** in the sidebar.
    - Choose a weakness area (e.g., *Fire Dynamics*).
    - The system will pull specific questions identified in your Study Plan, regardless of whether you have seen them before.
    """)

elif st.session_state.exam_stage == "ACTIVE":
    q_idx = st.session_state.current_index
    q = st.session_state.exam_session[q_idx]
    total_q = len(st.session_state.exam_session)

    col1, col2, col3 = st.columns([1, 4, 1])
    with col1:
        st.caption(f"Question {q_idx + 1} / {total_q}")
    with col3:
        flag_label = "🚩 Flagged" if q.flagged else "🏳️ Flag"
        if st.button(flag_label):
            q.flagged = not q.flagged
            st.rerun()

    st.progress((q_idx + 1) / total_q)
    st.subheader(q.text)

    is_multi = len(q.correct_indices) > 1
    if is_multi:
        st.info("Select all that apply.")
        for i, opt in enumerate(q.options):
            checked = i in q.user_selections
            if st.checkbox(opt, value=checked, key=f"q{q_idx}_opt{i}"):
                q.user_selections.add(i)
            else:
                q.user_selections.discard(i)
    else:
        current_selection = list(q.user_selections)[0] if q.user_selections else None
        idx_val = current_selection if current_selection is not None else 0
        selected_option = st.radio(
            "Select one:",
            q.options,
            index=idx_val if current_selection is not None else None,
            key=f"q{q_idx}_radio"
        )
        if selected_option:
            sel_index = q.options.index(selected_option)
            q.user_selections = {sel_index}

    st.write("---")

    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if q_idx > 0:
            if st.button("<< Previous"):
                st.session_state.current_index -= 1
                st.rerun()
    with c2:
        if st.button("Review All Questions"):
            st.session_state.exam_stage = "REVIEW"
            st.rerun()
    with c3:
        if q_idx < total_q - 1:
            if st.button("Next >>"):
                st.session_state.current_index += 1
                st.rerun()
        else:
            if st.button("Go to Review"):
                st.session_state.exam_stage = "REVIEW"
                st.rerun()

elif st.session_state.exam_stage == "REVIEW":
    st.title("Review Your Answers")
    cols = st.columns(5)
    for i, q in enumerate(st.session_state.exam_session):
        status = "⬛"
        if q.is_answered(): status = "🟩"
        if q.flagged: status = "🚩"

        label = f"Q{i + 1} {status}"
        with cols[i % 5]:
            if st.button(label, key=f"nav_{i}"):
                st.session_state.current_index = i
                st.session_state.exam_stage = "ACTIVE"
                st.rerun()

    st.write("---")

    if st.button("SUBMIT EXAM", type="primary"):
        # Update Scores and Wrong IDs List
        s_correct = 0
        s_total = len(st.session_state.exam_session)

        for q in st.session_state.exam_session:
            if q.is_correct():
                s_correct += 1
            else:
                st.session_state.wrong_ids.add(q.id)

        # ADD to the total, don't overwrite it
        st.session_state.cumulative_correct += s_correct
        st.session_state.cumulative_answered += s_total

        st.session_state.exam_stage = "REPORT"
        st.rerun()

elif st.session_state.exam_stage == "REPORT":
    st.title("Exam Results")
    score = sum(1 for q in st.session_state.exam_session if q.is_correct())
    total = len(st.session_state.exam_session)
    percent = (score / total) * 100

    st.metric("Session Score", f"{percent:.1f}%", f"{score}/{total} Correct")

    # 1. WAR REPORT BUTTON (For Immediate Review)
    wrong_qs = [q for q in st.session_state.exam_session if not q.is_correct()]
    if wrong_qs:
        war_report = generate_war_report(wrong_qs, title="Session Wrong Answer Roundup")
        st.download_button(
            label="⚠️ Download Session WAR Report",
            data=war_report,
            file_name="Session_WAR_Report.txt",
            mime="text/plain"
        )
    else:
        st.success("Perfect Score! No WAR Report needed.")

    # Results Table
    results_data = []
    for i, q in enumerate(st.session_state.exam_session):
        status = "✅ CORRECT" if q.is_correct() else "❌ INCORRECT"
        results_data.append({
            "Q#": i + 1,
            "ID": q.id,
            "Result": status,
            "Flagged": "Yes" if q.flagged else ""
        })
    st.table(results_data)

    if st.button("Start Next Exam Block"):
        st.session_state.exam_stage = "SETUP"
        st.rerun()